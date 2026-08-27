#!/usr/bin/env python3
"""Build styled bilingual V3 manuscript DOCX files from verified Markdown."""

from __future__ import annotations

import argparse
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PANDOC = Path("/opt/homebrew/bin/pandoc")
TABLE_INDENT_DXA = 120
ACCENT = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 108)
HEADER_FILL = "F4F6F9"


def _set_font(run: object, name: str, size: float | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def _configure_style(
    style: object,
    *,
    font: str,
    size: float,
    color: RGBColor,
    before: float,
    after: float,
    line_spacing: float,
    bold: bool = False,
) -> None:
    style.font.name = font
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font)
    paragraph = style.paragraph_format
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line_spacing
    paragraph.widow_control = True


def _page_field(paragraph: object) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, text, end))


def create_reference_docx(path: Path, *, language: str) -> None:
    """Create a narrative-proposal reference DOCX with academic overrides."""

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    body_font = "Source Han Serif CN" if language == "zh" else "Calibri"
    heading_font = "Source Han Serif CN" if language == "zh" else "Calibri"
    _configure_style(
        document.styles["Normal"],
        font=body_font,
        size=11,
        color=RGBColor(0, 0, 0),
        before=0,
        after=8,
        line_spacing=1.333,
    )
    _configure_style(
        document.styles["Title"],
        font=heading_font,
        size=22,
        color=RGBColor(11, 37, 69),
        before=0,
        after=12,
        line_spacing=1.05,
        bold=True,
    )
    _configure_style(
        document.styles["Subtitle"],
        font=body_font,
        size=10,
        color=MUTED,
        before=0,
        after=18,
        line_spacing=1.1,
    )
    for name, size, color, before, after in (
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ):
        _configure_style(
            document.styles[name],
            font=heading_font,
            size=size,
            color=color,
            before=before,
            after=after,
            line_spacing=1.1,
            bold=True,
        )
        document.styles[name].paragraph_format.keep_with_next = True
    _configure_style(
        document.styles["Caption"],
        font=body_font,
        size=9,
        color=MUTED,
        before=4,
        after=8,
        line_spacing=1.0,
    )
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run(
        "SR-SC-NARR | V3 Manuscript Draft"
        if language == "en"
        else "SR-SC-NARR | V3 中文论文初稿"
    )
    _set_font(header_run, heading_font, 8.5)
    header_run.font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Page " if language == "en" else "第 ")
    _set_font(footer_run, body_font, 8.5)
    footer_run.font.color.rgb = MUTED
    _page_field(footer)
    if language == "zh":
        ending = footer.add_run(" 页")
        _set_font(ending, body_font, 8.5)
        ending.font.color.rgb = MUTED
    document.save(path)


def _set_table_geometry(table: object, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        layout_anchor = properties.find(qn("w:tblLayout"))
        if layout_anchor is None:
            properties.append(indent)
        else:
            properties.insert(properties.index(layout_anchor), indent)
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            margins = cell_properties.find(qn("w:tcMar"))
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell_properties.append(margins)
            for side, value in (
                ("top", 80),
                ("bottom", 80),
                ("start", 120),
                ("end", 120),
            ):
                element = margins.find(qn(f"w:{side}"))
                if element is None:
                    element = OxmlElement(f"w:{side}")
                    margins.append(element)
                element.set(qn("w:w"), str(value))
                element.set(qn("w:type"), "dxa")


def _header_row(row: object) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)
    for cell in row.cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), HEADER_FILL)
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def postprocess_docx(path: Path, *, language: str) -> None:
    document = Document(path)
    body_font = "Source Han Serif CN" if language == "zh" else "Calibri"
    heading_font = "Source Han Serif CN" if language == "zh" else "Calibri"
    _configure_style(
        document.styles["Normal"],
        font=body_font,
        size=11,
        color=RGBColor(0, 0, 0),
        before=0,
        after=8,
        line_spacing=1.333,
    )
    _configure_style(
        document.styles["Title"],
        font=heading_font,
        size=22,
        color=RGBColor(11, 37, 69),
        before=0,
        after=12,
        line_spacing=1.05,
        bold=True,
    )
    if "Subtitle" in document.styles:
        _configure_style(
            document.styles["Subtitle"],
            font=body_font,
            size=10,
            color=MUTED,
            before=0,
            after=18,
            line_spacing=1.1,
        )
    for name, size, color, before, after in (
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ):
        _configure_style(
            document.styles[name],
            font=heading_font,
            size=size,
            color=color,
            before=before,
            after=after,
            line_spacing=1.1,
            bold=True,
        )
        document.styles[name].paragraph_format.keep_with_next = True
    _configure_style(
        document.styles["Caption"],
        font=body_font,
        size=9,
        color=MUTED,
        before=4,
        after=8,
        line_spacing=1.0,
    )
    document.core_properties.title = (
        "Spectral-Complexity-Gated Neural Augmentation"
        if language == "en"
        else "面向 Bloch 谱簇的谱复杂度门控神经增强方法"
    )
    document.core_properties.author = "Anonymous Authors"
    for paragraph in document.paragraphs:
        paragraph.paragraph_format.widow_control = True
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.text.startswith(("Figure ", "图")):
            paragraph.style = document.styles["Caption"]
            paragraph.paragraph_format.keep_with_next = True
        for run in paragraph.runs:
            _set_font(
                run,
                heading_font
                if paragraph.style.name.startswith(("Heading", "Title"))
                else body_font,
            )
    # Pandoc emits a fixed-layout academic table with repeated headers. Keep
    # that OOXML geometry intact and change typography only.
    for table in document.tables:
        for row in table.rows:
            row_properties = row._tr.get_or_add_trPr()
            if row_properties.find(qn("w:cantSplit")) is None:
                cant_split = OxmlElement("w:cantSplit")
                cant_split.set(qn("w:val"), "true")
                row_properties.append(cant_split)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        _set_font(run, body_font, 8.2)
    for shape in document.inline_shapes:
        if shape.width > Inches(6.25):
            ratio = Inches(6.25) / shape.width
            shape.width = Inches(6.25)
            shape.height = int(shape.height * ratio)
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        header = section.header.paragraphs[0]
        header._p.clear_content()
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_run = header.add_run(
            "SR-SC-NARR | V3 Manuscript Draft"
            if language == "en"
            else "SR-SC-NARR | V3 中文论文初稿"
        )
        _set_font(header_run, heading_font, 8.5)
        header_run.font.color.rgb = MUTED
        footer = section.footer.paragraphs[0]
        footer._p.clear_content()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer.add_run("Page " if language == "en" else "第 ")
        _set_font(footer_run, body_font, 8.5)
        footer_run.font.color.rgb = MUTED
        _page_field(footer)
        if language == "zh":
            ending = footer.add_run(" 页")
            _set_font(ending, body_font, 8.5)
            ending.font.color.rgb = MUTED
    document.save(path)


def build_one(
    markdown: Path,
    output: Path,
    *,
    language: str,
    root: Path,
) -> None:
    if not PANDOC.is_file():
        raise FileNotFoundError(f"pandoc is unavailable: {PANDOC}")
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="v3-docx-") as directory:
        normalized_markdown = Path(directory) / markdown.name
        source = markdown.read_text()
        source = re.sub(
            r"\\\((.+?)\\\)",
            lambda match: f"${match.group(1)}$",
            source,
            flags=re.DOTALL,
        )
        normalized_markdown.write_text(source)
        subprocess.run(
            (
                str(PANDOC),
                str(normalized_markdown),
                "--from=markdown+pipe_tables+fenced_code_blocks+tex_math_dollars+tex_math_single_backslash",
                "--to=docx",
                f"--resource-path={root}:{markdown.parent}",
                "--standalone",
                f"--output={output}",
            ),
            cwd=root,
            check=True,
        )
    postprocess_docx(output, language=language)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output-dir", type=Path, default=Path("paper/v3_submission")
    )
    args = parser.parse_args()
    root = Path(__file__).resolve().parents[1]
    build_one(
        root / "paper/v3_manuscript/MANUSCRIPT.en.md",
        args.output_dir / "SR-SC-NARR_manuscript_en.docx",
        language="en",
        root=root,
    )
    build_one(
        root / "paper/v3_manuscript/MANUSCRIPT.zh-CN.md",
        args.output_dir / "SR-SC-NARR_manuscript_zh-CN.docx",
        language="zh",
        root=root,
    )
    print(f"V3_DOCX_DIR={args.output_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
