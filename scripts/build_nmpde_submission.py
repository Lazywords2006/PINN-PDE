#!/usr/bin/env python3
"""Build editable NMPDE manuscript and cover-letter DOCX files."""

from __future__ import annotations

import argparse
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TITLE = (
    "A Basis-Invariant Neural-Augmented Rayleigh-Ritz Solver for Parametric "
    "Bloch Spectral Clusters with Eigenvalue Crossings"
)
JOURNAL = "Numerical Methods for Partial Differential Equations"
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(90, 90, 90)


def _font(
    run, *, size: float = 12, bold: bool | None = None, italic: bool | None = None
) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_font(
    style, *, size: float, bold: bool = False, italic: bool = False
) -> None:
    style.font.name = "Times New Roman"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = BLACK


def _page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))
    _font(run, size=10)


def _reference_doc(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    _set_style_font(normal, size=12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)

    title = styles["Title"]
    _set_style_font(title, size=15, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    title_properties = title._element.get_or_add_pPr()
    title_border = title_properties.find(qn("w:pBdr"))
    if title_border is not None:
        title_properties.remove(title_border)

    subtitle = styles["Subtitle"]
    _set_style_font(subtitle, size=12, italic=True)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)

    for name, size, bold, italic in (
        ("Heading 1", 12, True, False),
        ("Heading 2", 12, True, True),
        ("Heading 3", 12, False, True),
    ):
        style = styles[name]
        _set_style_font(style, size=size, bold=bold, italic=italic)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    _set_style_font(caption, size=10, italic=True)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_after = Pt(8)

    if "Author" not in styles:
        styles.add_style("Author", WD_STYLE_TYPE.PARAGRAPH)
    author = styles["Author"]
    _set_style_font(author, size=12)
    author.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.paragraph_format.line_spacing = 1.0
    author.paragraph_format.space_after = Pt(3)

    doc.add_paragraph("Reference document")
    doc.save(path)


def _clean_markdown(source: Path, destination: Path) -> None:
    lines = source.read_text().splitlines()
    start = next(index for index, line in enumerate(lines) if line == "## Abstract")
    body = "\n".join(lines[start:]) + "\n"
    metadata = (
        f'---\ntitle: "{TITLE}"\nauthor: "[Author Name(s)]"\nlang: en-US\n---\n\n'
    )
    destination.write_text(metadata + body)


def _insert_after(paragraph, text: str, *, italic: bool = False, size: float = 11):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    new_p.addnext(new_paragraph._p)
    run = new_paragraph.add_run(text)
    _font(run, size=size, italic=italic)
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    new_paragraph.paragraph_format.line_spacing = 1.0
    new_paragraph.paragraph_format.space_after = Pt(3)
    return new_paragraph


def _postprocess_manuscript(path: Path) -> None:
    doc = Document(path)
    for section in doc.sections:
        section.different_first_page_header_footer = False
        header = section.header.paragraphs[0]
        header.text = (
            "Original Article | Numerical Methods for Partial Differential Equations"
        )
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header.runs:
            _font(run, size=9, italic=True)
            run.font.color.rgb = MUTED
        footer = section.footer.paragraphs[0]
        footer.clear()
        _page_number(footer)

    author = next((p for p in doc.paragraphs if "[Author Name(s)]" in p.text), None)
    if author is not None:
        author.style = doc.styles["Author"]
        affiliation = _insert_after(
            author, "[Department, Institution, City, Country]", italic=True
        )
        corresponding = _insert_after(
            affiliation,
            "Corresponding author: [Name, email, ORCID]",
            italic=True,
            size=10.5,
        )
        _insert_after(
            corresponding, "Manuscript type: Original Article", italic=True, size=10.5
        )

    in_references = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name == "Title":
            properties = paragraph._element.get_or_add_pPr()
            border = properties.find(qn("w:pBdr"))
            if border is not None:
                properties.remove(border)
        if text == "References":
            in_references = True
        if text.startswith(("Table ", "Figure ")):
            paragraph.style = doc.styles["Caption"]
        if in_references and re.match(r"^\[\d+\]", text):
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(4)

    for table in doc.tables:
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        if len(table.columns) == 4:
            column_widths = (Inches(3.1), Inches(1.1), Inches(1.1), Inches(1.1))
        else:
            column_widths = (
                Inches(2.4),
                Inches(1.0),
                Inches(1.0),
                Inches(1.0),
                Inches(1.0),
            )
        for column, width in zip(table.columns, column_widths, strict=True):
            column.width = width
        header_row = table.rows[0]._tr
        tr_pr = header_row.get_or_add_trPr()
        repeat = OxmlElement("w:tblHeader")
        repeat.set(qn("w:val"), "true")
        tr_pr.append(repeat)
        for row_index, row in enumerate(table.rows):
            row_properties = row._tr.get_or_add_trPr()
            no_split = OxmlElement("w:cantSplit")
            row_properties.append(no_split)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                for paragraph in cell.paragraphs:
                    paragraph.style = doc.styles["Normal"]
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.keep_together = False
                    for run in paragraph.runs:
                        _font(run, size=9, bold=True if row_index == 0 else None)

    for shape in doc.inline_shapes:
        if shape.width > Inches(6.5):
            ratio = shape.height / shape.width
            shape.width = Inches(6.5)
            shape.height = int(shape.width * ratio)

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Original research manuscript"
    doc.core_properties.author = "[Author Name(s)]"
    doc.core_properties.keywords = (
        "neural PDE solver; parametric eigenproblem; Bloch-Schrodinger equation; "
        "spectral cluster; Rayleigh-Ritz; basis invariance; scientific machine learning"
    )
    doc.save(path)


def _cover_letter(source: Path, output: Path) -> None:
    text = re.sub(r"^# Cover Letter\s*", "", source.read_text(), count=1)
    blocks = [block.strip() for block in re.split(r"\n\s*\n", text) if block.strip()]
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    _set_style_font(normal, size=10.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(5)
    for block in blocks:
        if block.startswith("[Corresponding"):
            doc.add_paragraph()
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        line_breaks = "  \n" in block
        content = block.replace("  \n", "\n")
        if not line_breaks:
            content = re.sub(r"(?<!\n)\n(?!\n)", " ", content)
        content = content.replace("*", "")
        parts = content.split("\n")
        for index, part in enumerate(parts):
            if index:
                paragraph.add_run().add_break()
            run = paragraph.add_run(part.strip())
            _font(run, size=10.5, italic=part.strip() == JOURNAL)
    doc.core_properties.title = f"Cover letter - {TITLE}"
    doc.core_properties.author = "[Corresponding Author]"
    doc.save(output)


def build(root: Path, output_dir: Path) -> None:
    source = root / "paper/p2_final/MANUSCRIPT.en.md"
    cover_source = root / "paper/submission_nmpde/COVER_LETTER.en.md"
    output_dir.mkdir(parents=True, exist_ok=True)
    reference = output_dir / "NMPDE_reference.docx"
    manuscript = output_dir / "NMPDE_manuscript.docx"
    cover = output_dir / "NMPDE_cover_letter.docx"
    _reference_doc(reference)
    with tempfile.TemporaryDirectory(prefix="nmpde-") as temporary:
        cleaned = Path(temporary) / "manuscript.md"
        _clean_markdown(source, cleaned)
        subprocess.run(
            (
                "pandoc",
                str(cleaned),
                "--from=markdown+tex_math_dollars+tex_math_single_backslash+pipe_tables",
                "--to=docx",
                "--standalone",
                f"--reference-doc={reference}",
                f"--resource-path={root / 'paper/p2_final'}:{root}",
                f"--output={manuscript}",
            ),
            check=True,
        )
    _postprocess_manuscript(manuscript)
    _cover_letter(cover_source, cover)
    reference.unlink()
    print(f"MANUSCRIPT={manuscript}")
    print(f"COVER_LETTER={cover}")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output-dir", type=Path, default=Path("paper/submission_nmpde/final")
    )
    args = parser.parse_args()
    root = Path(__file__).resolve().parents[1]
    output = (
        args.output_dir if args.output_dir.is_absolute() else root / args.output_dir
    )
    build(root, output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
